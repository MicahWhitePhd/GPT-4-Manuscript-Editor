import docx
import os

class ManuscriptWriter:
    def __init__(self, file_path):
        self.file_path = file_path

    def add_paragraph(self, text):
        raise NotImplementedError

    def save(self):
        raise NotImplementedError

class DocxManuscriptWriter(ManuscriptWriter):
    def __init__(self, file_path):
        super().__init__(file_path)
        self.doc = docx.Document()

    def add_paragraph(self, text):
        self.doc.add_paragraph(text)

    def save(self):
        self.doc.save(self.file_path)

class TxtManuscriptWriter(ManuscriptWriter):
    def __init__(self, file_path):
        super().__init__(file_path)
        self.text = ""

    def add_paragraph(self, text):
        self.text += text + "\n"

    def save(self):
        with open(self.file_path, "w") as file:
            file.write(self.text)
